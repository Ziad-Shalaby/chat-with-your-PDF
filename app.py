import os
import streamlit as st
from pathlib import Path
from typing import List, Dict, Any
from io import BytesIO

# LangChain imports
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.document_loaders import PyPDFLoader, TextLoader
from langchain.docstore.document import Document
from langchain.embeddings import HuggingFaceEmbeddings
from langchain.vectorstores import FAISS
from langchain.chains import ConversationalRetrievalChain
from langchain.memory import ConversationBufferWindowMemory
from langchain.prompts import PromptTemplate
from langchain.llms import HuggingFaceHub

# Additional imports for file processing
from pypdf import PdfReader
import docx
import tempfile

# --------------------------- Config / Page ---------------------------
st.set_page_config(page_title="ChatWithYourDocs - LangChain", page_icon="📄", layout="wide")
st.title("📄💬 Chat-With-Your-Docs (LangChain Edition)")
st.caption("Upload multiple PDF / DOCX / TXT files, then chat with them using LangChain RAG pipeline.")

# --------------------------- Load Hugging Face Token ---------------------------
hf_token = st.secrets.get("HUGGINGFACEHUB_API_TOKEN", "")

# --------------------------- Document Processing ---------------------------
def process_uploaded_files(uploaded_files, chunk_size: int = 1000, chunk_overlap: int = 200) -> List[Document]:
    """Process uploaded files and return LangChain Documents"""
    documents = []
    
    for uploaded_file in uploaded_files:
        file_extension = Path(uploaded_file.name).suffix.lower()
        
        try:
            if file_extension == ".pdf":
                # Save uploaded file temporarily for PyPDFLoader
                with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp_file:
                    tmp_file.write(uploaded_file.read())
                    tmp_path = tmp_file.name
                
                # Use PyPDFLoader for better page handling
                loader = PyPDFLoader(tmp_path)
                docs = loader.load()
                
                # Add enhanced metadata
                for i, doc in enumerate(docs):
                    doc.metadata.update({
                        "source": uploaded_file.name,
                        "file_type": "pdf",
                        "page_number": i + 1,
                        "total_pages": len(docs)
                    })
                
                documents.extend(docs)
                os.unlink(tmp_path)  # Clean up temp file
                
            elif file_extension == ".txt":
                # Process TXT files
                content = uploaded_file.read().decode("utf-8", errors="ignore")
                
                # Create document with metadata
                doc = Document(
                    page_content=content,
                    metadata={
                        "source": uploaded_file.name,
                        "file_type": "txt",
                        "page_number": 1,
                        "total_pages": 1
                    }
                )
                documents.append(doc)
                
            elif file_extension == ".docx":
                # Process DOCX files
                file_bytes = uploaded_file.read()
                document = docx.Document(BytesIO(file_bytes))
                
                # Extract text from paragraphs
                paragraphs = []
                for paragraph in document.paragraphs:
                    if paragraph.text.strip():
                        paragraphs.append(paragraph.text)
                
                content = "\n\n".join(paragraphs)
                
                # Create document with metadata
                doc = Document(
                    page_content=content,
                    metadata={
                        "source": uploaded_file.name,
                        "file_type": "docx",
                        "page_number": 1,
                        "total_pages": 1,
                        "total_paragraphs": len(paragraphs)
                    }
                )
                documents.append(doc)
                
        except Exception as e:
            st.error(f"❌ Error processing {uploaded_file.name}: {e}")
            continue
    
    # Split documents into chunks using RecursiveCharacterTextSplitter
    if documents:
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            length_function=len,
            separators=["\n\n", "\n", ".", " ", ""]
        )
        
        split_docs = text_splitter.split_documents(documents)
        
        # Add chunk metadata
        for i, doc in enumerate(split_docs):
            doc.metadata["chunk_id"] = i
            doc.metadata["chunk_size"] = len(doc.page_content)
        
        return split_docs
    
    return []

# --------------------------- Vector Store Setup ---------------------------
@st.cache_resource(show_spinner=False)
def get_embeddings(model_name: str = "sentence-transformers/all-MiniLM-L6-v2"):
    """Get HuggingFace embeddings model with caching"""
    return HuggingFaceEmbeddings(
        model_name=model_name,
        model_kwargs={'device': 'cpu'},
        encode_kwargs={'normalize_embeddings': True}
    )

def create_vector_store(documents: List[Document], embeddings):
    """Create FAISS vector store from documents"""
    if not documents:
        raise ValueError("No documents to create vector store")
    
    # Create FAISS vector store
    vector_store = FAISS.from_documents(documents, embeddings)
    return vector_store

# --------------------------- LLM Setup ---------------------------
def get_llm(model_name: str, hf_token: str, max_new_tokens: int = 300, temperature: float = 0.7):
    """Get HuggingFace LLM"""
    return HuggingFaceHub(
        repo_id=model_name,
        huggingfacehub_api_token=hf_token,
        model_kwargs={
            "max_new_tokens": max_new_tokens,
            "temperature": temperature,
            "return_full_text": False
        }
    )

# --------------------------- Custom Prompt Template ---------------------------
def get_custom_prompt():
    """Create custom prompt template for RAG"""
    template = """You are a helpful AI assistant that answers questions based on the provided document context.

INSTRUCTIONS:
- Use ONLY the information from the provided context to answer questions
- If you cannot find the answer in the context, clearly state that you don't have that information
- Be concise but comprehensive in your responses
- Cite specific parts of the documents when relevant

CONTEXT FROM DOCUMENTS:
{context}

CHAT HISTORY:
{chat_history}

HUMAN QUESTION: {question}

ASSISTANT ANSWER:"""
    
    return PromptTemplate(
        template=template,
        input_variables=["context", "chat_history", "question"]
    )

# --------------------------- RAG Chain Setup ---------------------------
def create_rag_chain(vector_store, llm, custom_prompt, top_k: int = 3):
    """Create conversational retrieval chain"""
    # Memory to maintain conversation context
    memory = ConversationBufferWindowMemory(
        memory_key="chat_history",
        output_key="answer",
        return_messages=True,
        k=5  # Keep last 5 conversation turns
    )
    
    # Create the retrieval chain
    retrieval_chain = ConversationalRetrievalChain.from_llm(
        llm=llm,
        retriever=vector_store.as_retriever(
            search_type="similarity",
            search_kwargs={"k": top_k}
        ),
        memory=memory,
        combine_docs_chain_kwargs={"prompt": custom_prompt},
        return_source_documents=True,
        verbose=False
    )
    
    return retrieval_chain

# --------------------------- Session State Management ---------------------------
if "chat_history" not in st.session_state:
    st.session_state.chat_history = []
if "rag_chain" not in st.session_state:
    st.session_state.rag_chain = None
if "vector_store" not in st.session_state:
    st.session_state.vector_store = None
if "documents_ready" not in st.session_state:
    st.session_state.documents_ready = False
if "processed_files" not in st.session_state:
    st.session_state.processed_files = []
if "total_chunks" not in st.session_state:
    st.session_state.total_chunks = 0

# --------------------------- Sidebar Configuration ---------------------------
with st.sidebar:
    st.header("⚙️ Configuration")
    
    # Model Selection
    st.subheader("🤖 Models")
    model_name = st.selectbox(
        "Chat Model", 
        [
            "mistralai/Mistral-7B-Instruct-v0.3",
            "microsoft/DialoGPT-medium",
            "google/flan-t5-large",
            "HuggingFaceH4/zephyr-7b-beta",
            "meta-llama/Llama-2-7b-chat-hf"
        ], 
        index=0,
        help="Choose the language model for generating responses"
    )
    
    embedding_model = st.selectbox(
        "Embedding Model",
        [
            "sentence-transformers/all-MiniLM-L6-v2",
            "sentence-transformers/all-mpnet-base-v2",
            "sentence-transformers/paraphrase-multilingual-MiniLM-L12-v2",
            "sentence-transformers/multi-qa-MiniLM-L6-cos-v1"
        ],
        index=0,
        help="Choose the model for creating document embeddings"
    )
    
    # Text Processing Settings
    st.subheader("📝 Text Processing")
    chunk_size = st.slider(
        "Chunk Size (characters)", 
        500, 2000, 1000, 100,
        help="Size of text chunks for processing"
    )
    chunk_overlap = st.slider(
        "Chunk Overlap (characters)", 
        50, 500, 200, 50,
        help="Overlap between consecutive chunks"
    )
    
    # Retrieval & Generation Settings
    st.subheader("🔍 Retrieval & Generation")
    top_k = st.slider(
        "Top-K Retrieved Chunks", 
        1, 10, 3, 1,
        help="Number of most relevant chunks to retrieve"
    )
    max_new_tokens = st.slider(
        "Max New Tokens", 
        64, 1024, 300, 16,
        help="Maximum tokens in generated response"
    )
    temperature = st.slider(
        "Temperature", 
        0.0, 1.5, 0.7, 0.1,
        help="Controls randomness in responses (0=deterministic, 1.5=very creative)"
    )
    
    # Display Options
    st.subheader("📊 Display Options")
    show_sources = st.checkbox("Show Retrieved Sources", value=True)
    show_metadata = st.checkbox("Show Detailed Metadata", value=False)
    
    # Action Buttons
    st.subheader("🔄 Actions")
    col1, col2 = st.columns(2)
    
    with col1:
        if st.button("🧹 Clear Chat", use_container_width=True):
            st.session_state.chat_history = []
            if st.session_state.rag_chain and hasattr(st.session_state.rag_chain, 'memory'):
                st.session_state.rag_chain.memory.clear()
            st.success("Chat cleared!")
            st.rerun()
    
    with col2:
        if st.button("🗂️ Clear Docs", use_container_width=True):
            st.session_state.documents_ready = False
            st.session_state.rag_chain = None
            st.session_state.vector_store = None
            st.session_state.processed_files = []
            st.session_state.total_chunks = 0
            st.session_state.chat_history = []
            st.success("Documents cleared!")
            st.rerun()
    
    # Export Chat
    if st.session_state.chat_history:
        st.subheader("💾 Export")
        chat_transcript = "\n".join([
            f"{m['role'].upper()}: {m['content']}" 
            for m in st.session_state.chat_history
        ])
        
        st.download_button(
            "📥 Download Chat Transcript",
            data=chat_transcript,
            file_name="chat_transcript.txt",
            mime="text/plain",
            use_container_width=True
        )

# --------------------------- Main Content Area ---------------------------

# Document Upload Section
st.header("📁 Document Upload & Processing")

# Show current status
if st.session_state.processed_files:
    with st.expander("📋 Currently Loaded Documents", expanded=False):
        for file_info in st.session_state.processed_files:
            col1, col2, col3 = st.columns([3, 1, 1])
            with col1:
                st.write(f"📄 **{file_info['name']}**")
            with col2:
                st.write(f"`{file_info['type']}`")
            with col3:
                st.write(f"{file_info['chunks']} chunks")

# File uploader
uploaded_files = st.file_uploader(
    "Choose files to upload",
    type=["pdf", "docx", "txt"],
    accept_multiple_files=True,
    help="Upload PDF, DOCX, or TXT files. Multiple files are supported."
)

if uploaded_files:
    # Check if reprocessing is needed
    current_file_names = [f.name for f in uploaded_files]
    processed_file_names = [f['name'] for f in st.session_state.processed_files]
    
    needs_processing = (
        set(current_file_names) != set(processed_file_names) or
        not st.session_state.documents_ready
    )
    
    if needs_processing:
        with st.spinner("🔄 Processing documents... This may take a moment."):
            try:
                # Process documents
                documents = process_uploaded_files(uploaded_files, chunk_size, chunk_overlap)
                
                if documents:
                    # Create embeddings
                    embeddings = get_embeddings(embedding_model)
                    
                    # Create vector store
                    vector_store = create_vector_store(documents, embeddings)
                    st.session_state.vector_store = vector_store
                    
                    # Create RAG chain if token is available
                    if hf_token:
                        try:
                            llm = get_llm(model_name, hf_token, max_new_tokens, temperature)
                            custom_prompt = get_custom_prompt()
                            rag_chain = create_rag_chain(vector_store, llm, custom_prompt, top_k)
                            
                            st.session_state.rag_chain = rag_chain
                            st.session_state.documents_ready = True
                            st.session_state.total_chunks = len(documents)
                            
                            # Update processed files tracking
                            file_info = {}
                            for doc in documents:
                                source = doc.metadata.get("source", "Unknown")
                                file_type = doc.metadata.get("file_type", "unknown")
                                
                                if source not in file_info:
                                    file_info[source] = {"type": file_type.upper(), "chunks": 0}
                                file_info[source]["chunks"] += 1
                            
                            st.session_state.processed_files = [
                                {"name": source, "type": info["type"], "chunks": info["chunks"]}
                                for source, info in file_info.items()
                            ]
                            
                            # Success message
                            st.success(f"✅ Successfully processed {len(documents)} chunks from {len(uploaded_files)} documents!")
                            
                            # Display statistics
                            col1, col2, col3, col4 = st.columns(4)
                            with col1:
                                st.metric("📄 Documents", len(uploaded_files))
                            with col2:
                                st.metric("🧩 Chunks", len(documents))
                            with col3:
                                avg_chunk_size = sum(len(doc.page_content) for doc in documents) // len(documents)
                                st.metric("📏 Avg Chunk Size", f"{avg_chunk_size} chars")
                            with col4:
                                total_chars = sum(len(doc.page_content) for doc in documents)
                                st.metric("📊 Total Content", f"{total_chars:,} chars")
                        
                        except Exception as llm_error:
                            st.error(f"❌ Error setting up language model: {str(llm_error)}")
                            st.info("Try selecting a different model or check your Hugging Face token permissions.")
                    
                    else:
                        st.error("❌ Hugging Face token not found in Streamlit secrets!")
                        st.info("Please add HUGGINGFACEHUB_API_TOKEN to your Streamlit secrets.")
                
                else:
                    st.warning("⚠️ No content could be extracted from the uploaded files.")
                    
            except Exception as e:
                st.error(f"❌ Error processing documents: {str(e)}")
                st.error("Please check your files and try again.")

# Chat Interface Section
st.header("💬 Chat with Your Documents")

if not st.session_state.documents_ready:
    st.info("👆 Please upload and process documents above to start chatting!")
elif not hf_token:
    st.error("❌ Hugging Face API token not found. Please check your Streamlit secrets configuration.")
else:
    # Display chat messages
    for message in st.session_state.chat_history:
        with st.chat_message(message["role"]):
            st.write(message["content"])
    
    # Chat input
    if user_question := st.chat_input("Ask me anything about your documents..."):
        # Add user message to history
        st.session_state.chat_history.append({"role": "user", "content": user_question})
        
        # Display user message
        with st.chat_message("user"):
            st.write(user_question)
        
        # Generate and display response
        with st.chat_message("assistant"):
            with st.spinner("🤔 Thinking..."):
                try:
                    # Get response from RAG chain
                    result = st.session_state.rag_chain({
                        "question": user_question
                    })
                    
                    response = result.get("answer", "No answer generated.")
                    source_documents = result.get("source_documents", [])
                    
                    # Clean up response if needed
                    if isinstance(response, str) and response.strip():
                        response = response.strip()
                    else:
                        response = "I couldn't generate a proper response. Please try rephrasing your question."
                    
                    # Display response
                    st.write(response)
                    
                    # Display sources if enabled
                    if show_sources and source_documents:
                        with st.expander(f"🔍 Sources ({len(source_documents)} documents used)"):
                            for i, doc in enumerate(source_documents, 1):
                                source = doc.metadata.get("source", "Unknown")
                                page = doc.metadata.get("page", doc.metadata.get("page_number", "N/A"))
                                file_type = doc.metadata.get("file_type", "unknown")
                                
                                st.markdown(f"**Source {i}: {source}**")
                                
                                if show_metadata:
                                    col1, col2 = st.columns(2)
                                    with col1:
                                        st.write(f"📄 Type: {file_type.upper()}")
                                        st.write(f"📖 Page: {page}")
                                    with col2:
                                        chunk_id = doc.metadata.get("chunk_id", "N/A")
                                        chunk_size = doc.metadata.get("chunk_size", len(doc.page_content))
                                        st.write(f"🧩 Chunk ID: {chunk_id}")
                                        st.write(f"📏 Size: {chunk_size} chars")
                                
                                # Show content preview
                                content = doc.page_content
                                if len(content) > 300:
                                    content = content[:300] + "..."
                                
                                st.markdown(f"```\n{content}\n```")
                                
                                if i < len(source_documents):
                                    st.divider()
                    
                    # Add response to history
                    st.session_state.chat_history.append({"role": "assistant", "content": response})
                    
                except Exception as e:
                    error_msg = f"❌ I encountered an error: {str(e)[:200]}..."
                    st.error("I'm having trouble generating a response. This could be due to:")
                    st.error("- Model server being busy")
                    st.error("- Token rate limits")
                    st.error("- Network connectivity issues")
                    st.info("💡 Try: Selecting a different model, waiting a moment, or rephrasing your question.")
                    
                    # Add error to history for context
                    st.session_state.chat_history.append({
                        "role": "assistant", 
                        "content": "I apologize, but I encountered a technical issue. Please try again."
                    })

# Footer and Help Section
st.markdown("---")

col1, col2 = st.columns([3, 1])

with col1:
    st.markdown("**🦜 Built with LangChain & Streamlit**")
    if st.session_state.documents_ready:
        st.success(f"📊 System Status: {len(st.session_state.processed_files)} documents loaded, {st.session_state.total_chunks} chunks indexed")
    else:
        st.info("📊 System Status: Ready for document upload")

with col2:
    if st.button("❓ Show Help"):
        st.info("Expand the help section below!")

# Help section
with st.expander("❓ How to Use This App"):
    st.markdown("""
    ### 🚀 Quick Start Guide:
    
    1. **Configure Settings** 📊
       - Choose your preferred models in the sidebar
       - Adjust chunk size and overlap for your content type
       - Set retrieval and generation parameters
    
    2. **Upload Documents** 📁
       - Click "Browse files" and select PDF, DOCX, or TXT files
       - Multiple files can be uploaded simultaneously
       - Wait for processing to complete
    
    3. **Start Chatting** 💬
       - Ask questions about your uploaded documents
       - View retrieved sources to see which parts were used
       - Export your conversation history if needed
    
    ### 💡 Pro Tips:
    
    - **Chunk Size**: Larger chunks provide more context, smaller chunks are more precise
    - **Temperature**: Lower values (0.1-0.3) for factual responses, higher (0.7-1.0) for creative ones
    - **Top-K**: More retrieved chunks give broader context but may include less relevant info
    - **Sources**: Enable source display to verify which document sections inform responses
    
    ### 🔧 Troubleshooting:
    
    - **No Response**: Check if Hugging Face token is properly set in Streamlit secrets
    - **Processing Errors**: Try smaller files or check file formats (PDF, DOCX, TXT only)
    - **Generic Responses**: Ensure your questions relate to uploaded document content
    - **Slow Performance**: Try reducing chunk size or number of retrieved documents
    
    ### 📚 Supported File Types:
    - **PDF**: Multi-page documents with text extraction
    - **DOCX**: Microsoft Word documents 
    - **TXT**: Plain text files with UTF-8 encoding
    """)

# Current configuration display
if st.session_state.documents_ready:
    with st.expander("⚙️ Current System Configuration"):
        config_col1, config_col2, config_col3 = st.columns(3)
        
        with config_col1:
            st.markdown("**🤖 AI Models**")
            st.write(f"Chat: `{model_name.split('/')[-1]}`")
            st.write(f"Embeddings: `{embedding_model.split('/')[-1]}`")
        
        with config_col2:
            st.markdown("**📝 Processing**")
            st.write(f"Chunk Size: {chunk_size:,} chars")
            st.write(f"Overlap: {chunk_overlap:,} chars")
        
        with config_col3:
            st.markdown("**🔍 Generation**")
            st.write(f"Top-K: {top_k} docs")
            st.write(f"Max Tokens: {max_new_tokens:,}")
            st.write(f"Temperature: {temperature}")
